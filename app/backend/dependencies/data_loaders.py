''' Loaders  '''

"""Connectors to to ingest external data"""

import os
from os import listdir
from os.path import isfile, join
import logging
import docx
from langchain.document_loaders import  WebBaseLoader, OnlinePDFLoader
from langchain.document_loaders import PyPDFLoader
from langchain.document_loaders import TextLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain.docstore.document import Document
import requests
import io
from typing import Union, Optional, Any, List
from urllib.request import urlopen
from io import BytesIO
import fitz
from fastapi import UploadFile

class PyMuPDFLoader():
    """Loader that uses PyMuPDF to load PDF files."""

    def __init__(self, file_source: Union[str, io.BytesIO], file_name: Optional[str] = None):
        """Initialize with file path or file-like object."""
        try:
            import fitz  # noqa:F401
        except ImportError:
            raise ValueError(
                "PyMuPDF package not found, please install it with "
                "`pip install pymupdf`"
            )
        self.file_source = file_source
        self.file_name = file_name

    def load(self, **kwargs: Optional[Any]) -> List[Document]:
        """Load file."""
        import fitz

        if isinstance(self.file_source, str):
            # File path provided
            file_path = self.file_source
            doc = fitz.open(filename=file_path)  # open document
        else:
            # File-like object provided
            file_path = None
            doc = fitz.open(stream=self.file_source, filetype="pdf")  # open document

        return [
            Document(
                page_content=page.get_text(**kwargs).encode("utf-8"),
                metadata=dict(
                    {
                        "source": self.file_name,
                        "file_path": self.file_name,
                        "page_number": page.number + 1,
                        "total_pages": len(doc),
                    },
                    **{
                        k: doc.metadata[k]
                        for k in doc.metadata
                        if type(doc.metadata[k]) in [str, int]
                        if doc.metadata[k] is not None
                    }
                ),
            )
            for page in doc
        ]



def split_docs(docs):
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
    return text_splitter.split_documents(docs)

def split_text(text):
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
    text_list = text_splitter.split_text(text)
    return text_splitter.create_documents(text_list)


def file_loader(file: UploadFile):
    '''Parses any type of data'''

    content = file.file.read()
    file_name = file.filename
    docs = []

    if file.content_type == 'application/pdf':
        docs = PyMuPDFLoader(BytesIO(content), file_name=file_name).load()
    elif file.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        doc = docx.Document(BytesIO(content))
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        docs = split_text('\n'.join(text))
    elif file.content_type == 'text/plain':
        loader = TextLoader(BytesIO(content))
        documents = loader.load()
        docs = split_docs(documents)
    else:
        print("The document was not the right format")

    logging.info(f"Loaded {len(docs)} documents")
    return docs


def url_loader(url):
    '''Parses web page'''
    if '.pdf' in url: 
        docs = OnlinePDFLoader(url.load())
    else:
        try:
            docs = WebBaseLoader(url).load()
        except requests.exceptions.MissingSchema:
            url = 'http://' + url
            docs = WebBaseLoader(url).load()
        except Exception as e:
            docs = []
    return docs 




